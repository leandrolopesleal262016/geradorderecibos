from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import pandas as pd
from datetime import datetime
import os
import zipfile
import locale
from datetime import datetime
from models import db, ReceiptSequence, ReciboGerado, ModeloRecibo, Cliente
from werkzeug.utils import secure_filename
from flask import abort
import json
from PIL import Image

# Create Flask app first
app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///recibos.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)

UPLOAD_FOLDER = 'static/images/logos'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Criar pasta de upload se não existir
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)


# Define helper function
def get_document_content(blob):
    doc = Document(io.BytesIO(blob))
    conteudo = []
    for paragrafo in doc.paragraphs:
        if paragrafo.text.strip():
            conteudo.append(paragrafo.text.strip())
    return conteudo

# Register function in Jinja environment
app.jinja_env.globals.update(get_document_content=get_document_content)

data_atual = datetime.now().strftime('%d/%m/%Y')
fornecedores_df = None
documentos_gerados = []  # Declaração global
def traduzir_mes(mes_en):
    meses = {
        'January': 'janeiro',
        'February': 'fevereiro',
        'March': 'março',
        'April': 'abril',
        'May': 'maio',
        'June': 'junho',
        'July': 'julho',
        'August': 'agosto',
        'September': 'setembro',
        'October': 'outubro',        'November': 'novembro',
        'December': 'dezembro'
    }
    return meses.get(mes_en, mes_en)

def formatar_data_atual():
    data_atual = datetime.now()
    mes_en = data_atual.strftime('%B')
    mes_pt = traduzir_mes(mes_en)
    return data_atual.strftime(f'%d de {mes_pt} de %Y')

def processar_modelo(modelo_conteudo, dados_cliente):
    texto = modelo_conteudo.replace('{cliente_nome}', dados_cliente['nome'])
    texto = texto.replace('{valor}', dados_cliente['valor'])
    texto = texto.replace('{valor_extenso}', dados_cliente['valor_extenso'])
    texto = texto.replace('{numero_documento}', dados_cliente['numero_documento'])
    texto = texto.replace('{data}', dados_cliente['data'])
    return texto.split('\n')

def valor_por_extenso(valor):
    # Converte para float caso seja string
    if isinstance(valor, str):
        valor = float(valor.replace('.', '').replace(',', '.'))
    
    # Separa parte inteira e decimal
    parte_inteira = int(valor)
    parte_decimal = int(round((valor - parte_inteira) * 100))

    # Arrays com palavras
    unidades = ['', 'um', 'dois', 'três', 'quatro', 'cinco', 'seis', 'sete', 'oito', 'nove']
    dezenas = ['', 'dez', 'vinte', 'trinta', 'quarenta', 'cinquenta', 'sessenta', 'setenta', 'oitenta', 'noventa']
    dezenas_especiais = ['dez', 'onze', 'doze', 'treze', 'quatorze', 'quinze', 'dezesseis', 'dezessete', 'dezoito', 'dezenove']
    centenas = ['', 'cento', 'duzentos', 'trezentos', 'quatrocentos', 'quinhentos', 'seiscentos', 'setecentos', 'oitocentos', 'novecentos']

    extenso = []

    # Trata valor zero
    if parte_inteira == 0 and parte_decimal == 0:
        return 'zero reais'

    # Processa milhares
    if parte_inteira >= 1000:
        milhares = parte_inteira // 1000
        if milhares == 1:
            extenso.append('um mil')
        else:
            if milhares >= 100:
                centena_milhar = milhares // 100
                extenso.append(centenas[centena_milhar])
                milhares = milhares % 100
                if milhares > 0:
                    extenso.append('e')
            
            if milhares >= 10:
                if milhares >= 10 and milhares <= 19:
                    extenso.append(dezenas_especiais[milhares - 10])
                else:
                    dezena_milhar = milhares // 10
                    extenso.append(dezenas[dezena_milhar])
                    milhares = milhares % 10
                    if milhares > 0:
                        extenso.append('e')
                        extenso.append(unidades[milhares])
            elif milhares > 0:
                extenso.append(unidades[milhares])
            
            extenso.append('mil')

        parte_inteira = parte_inteira % 1000
        if parte_inteira > 0:
            if parte_inteira < 100:
                extenso.append('e')
            else:
                extenso.append('')

    # Processa centenas
    if parte_inteira >= 100:
        centena = parte_inteira // 100
        if centena == 1 and parte_inteira % 100 == 0:
            extenso.append('cem')
        else:
            extenso.append(centenas[centena])
        parte_inteira = parte_inteira % 100
        if parte_inteira > 0:
            extenso.append('e')

    # Processa dezenas e unidades
    if parte_inteira >= 10 and parte_inteira <= 19:
        extenso.append(dezenas_especiais[parte_inteira - 10])
    else:
        if parte_inteira >= 10:
            extenso.append(dezenas[parte_inteira // 10])
            parte_inteira = parte_inteira % 10
            if parte_inteira > 0:
                extenso.append('e')
        if parte_inteira > 0:
            extenso.append(unidades[parte_inteira])

    extenso.append('reais')

    # Processa centavos
    if parte_decimal > 0:
        extenso.append('e')
        if parte_decimal >= 10 and parte_decimal <= 19:
            extenso.append(dezenas_especiais[parte_decimal - 10])
        else:
            dezena = parte_decimal // 10
            unidade = parte_decimal % 10
            if dezena > 0:
                extenso.append(dezenas[dezena])
                if unidade > 0:
                    extenso.append('e')
                    extenso.append(unidades[unidade])
            elif unidade > 0:
                extenso.append(unidades[unidade])
        extenso.append('centavos')

    return ' '.join(extenso)
def numero_para_extenso(numero):
    unidades = ['', 'um', 'dois', 'três', 'quatro', 'cinco', 'seis', 'sete', 'oito', 'nove']
    dezenas = ['', 'dez', 'vinte', 'trinta', 'quarenta', 'cinquenta', 'sessenta', 'setenta', 'oitenta', 'noventa']
    
    extenso = []
    
    # Verifica se o número está dentro do intervalo válido
    if numero < 0 or numero > 99:
        return ['número fora do intervalo válido']
    
    if numero >= 10:
        dezena = int(numero / 10)
        if dezena < len(dezenas):  # Verifica se o índice é válido
            extenso.append(dezenas[dezena])
            numero = numero % 10
            if numero > 0:
                extenso.append('e')
    
    if numero > 0 and numero < len(unidades):
        extenso.append(unidades[numero])
    
    return extenso

def get_next_receipt_number():
    with app.app_context():
        seq = ReceiptSequence.query.first()
        if not seq:
            seq = ReceiptSequence(last_number=0)
            db.session.add(seq)
        
        seq.last_number += 1
        db.session.commit()
        return f"{seq.last_number:05d}"
    
def is_cpf(documento):
    # Remove caracteres não numéricos
    doc = ''.join(filter(str.isdigit, documento))
    # CPF normalmente tem 11 dígitos, mas vamos considerar uma margem
    return len(doc) <= 12

def is_cnpj(documento):
    # Remove caracteres não numéricos
    doc = ''.join(filter(str.isdigit, documento))
    # CNPJ normalmente tem 14 dígitos
    return len(doc) > 12

def validar_cpf(cpf):
    # Remove caracteres não numéricos
    cpf = ''.join(filter(str.isdigit, cpf))
    
    # Verifica se tem 11 dígitos
    if len(cpf) != 11:
        return False
        
    # Verifica se todos os dígitos são iguais
    if len(set(cpf)) == 1:
        return False
        
    # Validação do primeiro dígito verificador
    soma = sum(int(cpf[i]) * (10 - i) for i in range(9))
    digito = (soma * 10) % 11
    if digito == 10:
        digito = 0
    if int(cpf[9]) != digito:
        return False
        
    # Validação do segundo dígito verificador
    soma = sum(int(cpf[i]) * (11 - i) for i in range(10))
    digito = (soma * 10) % 11
    if digito == 10:
        digito = 0
    if int(cpf[10]) != digito:
        return False
        
    return True

def validar_cnpj(cnpj):
    # Remove caracteres não numéricos
    cnpj = ''.join(filter(str.isdigit, cnpj))
    
    # Verifica se tem 14 dígitos
    if len(cnpj) != 14:
        return False
        
    # Verifica se todos os dígitos são iguais
    if len(set(cnpj)) == 1:
        return False
        
    # Validação do primeiro dígito verificador
    multiplicadores = [5,4,3,2,9,8,7,6,5,4,3,2]
    soma = sum(int(cnpj[i]) * multiplicadores[i] for i in range(12))
    digito = soma % 11
    if digito < 2:
        digito = 0
    else:
        digito = 11 - digito
    if int(cnpj[12]) != digito:
        return False
        
    # Validação do segundo dígito verificador
    multiplicadores = [6,5,4,3,2,9,8,7,6,5,4,3,2]
    soma = sum(int(cnpj[i]) * multiplicadores[i] for i in range(13))
    digito = soma % 11
    if digito < 2:
        digito = 0
    else:
        digito = 11 - digito
    if int(cnpj[13]) != digito:
        return False
        
    return True

    
@app.route('/')
def index():
    return render_template('index.html')

def is_cpf(documento):
    # Remove caracteres não numéricos
    doc = ''.join(filter(str.isdigit, documento))
    # CPF normalmente tem 11 dígitos, mas vamos considerar uma margem
    return len(doc) <= 11

def is_cnpj(documento):
    # Remove caracteres não numéricos
    doc = ''.join(filter(str.isdigit, documento))
    # CNPJ normalmente tem 14 dígitos
    return len(doc) > 11

@app.route('/get_clientes', methods=['GET'])
def get_clientes():
    try:
        # Busca todos os clientes
        todos_clientes = Cliente.query.order_by(Cliente.razao_social).all()
        
        empresas = []
        pessoas = []
        
        # Garante que todo cliente seja classificado
        for cliente in todos_clientes:
            if is_cnpj(cliente.cpf_cnpj):
                empresas.append(cliente.razao_social)
            else:
                # Se não for CNPJ, considera como CPF
                pessoas.append(cliente.razao_social)
        
        response_data = {
            'empresas': empresas,
            'pessoas': pessoas
        }
        
        print(f"Total de empresas: {len(empresas)}")
        print(f"Total de pessoas: {len(pessoas)}")
        print(f"Total de clientes: {len(todos_clientes)}")
        
        return jsonify(response_data)
        
    except Exception as e:
        print(f"Erro ao buscar clientes: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/generate_receipts_bulk', methods=['POST'])
def generate_receipts_bulk():
    global documentos_gerados
    try:
        dados = request.json
        modelo_id = dados.get('modelo')
        clientes_selecionados = dados.get('clientes', [])
        
        # Recebe a data do frontend e converte para objeto datetime
        data_enviada = dados.get('data')
        print(f"Data recebida do frontend: {data_enviada}")
        
        if data_enviada:
            data_selecionada = datetime.strptime(data_enviada, '%Y-%m-%d')
            print(f"Data convertida: {data_selecionada}")
        else:
            data_selecionada = datetime.now()
            print("Nenhuma data recebida, usando data atual.")

        # Formata a data para exibição no recibo (dia/mês/ano)
        data_formatada = data_selecionada.strftime('%d/%m/%Y')
        print(f"Data formatada para o texto: {data_formatada}")

        # Busca modelo no banco
        modelo = ModeloRecibo.query.get(modelo_id)
        if not modelo:
            return jsonify({'erro': 'Modelo não encontrado'}), 404
            
        print(f"Usando modelo {modelo_id}: {modelo.nome}")
        modelo_texto = modelo.conteudo

        valor_str = dados.get('valor', '0,00')
        valor_limpo = valor_str.replace('.', '').replace(',', '.')
        valor_float = float(valor_limpo)
        valor_formatado = f"{valor_float:,.2f}".replace(',', '_').replace('.', ',').replace('_', '.')

        documentos_gerados = []
        preview_content = []

        for cliente_nome in clientes_selecionados:
            numero_recibo = get_next_receipt_number()

            # Busca cliente no banco de dados
            cliente = Cliente.query.filter_by(razao_social=cliente_nome).first()
            if not cliente:
                continue

            texto_formatado = modelo_texto.format(
                cliente_nome=cliente.razao_social,
                valor=valor_formatado,
                valor_extenso=valor_por_extenso(valor_float),
                numero_recibo=numero_recibo,
                data=data_formatada,
                documento_cliente=cliente.cpf_cnpj
            )

            # Criação do documento Word
            doc = Document()
            sections = doc.sections
            for section in sections:
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Recebe a preferência do frontend
            mostrar_logo = dados.get('mostrarLogo', True)

            # Verifica se deve mostrar o logo usando o valor recebido do frontend
            if mostrar_logo:
                # Cabeçalho com logo e informações
                header_table = doc.add_table(rows=1, cols=2)
                header_table.autofit = False
                header_table.columns[0].width = Inches(1.2)
                header_table.columns[1].width = Inches(5.8)
                
                # Logo
                logo_cell = header_table.cell(0, 0)
                logo_paragraph = logo_cell.paragraphs[0]
                logo_run = logo_paragraph.add_run()
                
                # Adicionar logo
                if modelo.logo_path:
                    logo_path = modelo.logo_path.lstrip('/')  # Remove a barra inicial se existir
                    if os.path.exists(logo_path):
                        try:
                            logo_run.add_picture(logo_path, width=Inches(1.2))
                        except Exception as e:
                            print(f"Erro ao adicionar logo personalizada: {str(e)}")
                            logo_run.add_picture('static/images/logo.png', width=Inches(1.2))
                    else:
                        logo_run.add_picture('static/images/logo.png', width=Inches(1.2))
                else:
                    logo_run.add_picture('static/images/logo.png', width=Inches(1.2))
                
                # Texto do cabeçalho na coluna 2
                info_cell = header_table.cell(0, 1)
                info_paragraph = info_cell.paragraphs[0]
                header_text = modelo.header_text if modelo.header_text else "BEIJO E MATOS CONSTRUÇÕES E ENGENHARIA LTDA\nJoaquim da Silva Martha, 12-53 - Sala 3 - Altos da Cidade - Bauru/SP\nguilhermebeijo@bencato.com.br - CNPJ: 26.149.105/0001-09 - www.bencato.com.br"
                info_run = info_paragraph.add_run(header_text)
                info_run.font.color.rgb = RGBColor(128, 128, 128)
                info_run.font.size = Pt(11)
            else:
                # Se não mostrar logo, adicione apenas o texto do cabeçalho em um parágrafo
                header_paragraph = doc.add_paragraph()
                header_text = modelo.header_text if modelo.header_text else "BEIJO E MATOS CONSTRUÇÕES E ENGENHARIA LTDA\nJoaquim da Silva Martha, 12-53 - Sala 3 - Altos da Cidade - Bauru/SP\nguilhermebeijo@bencato.com.br - CNPJ: 26.149.105/0001-09 - www.bencato.com.br"
                header_run = header_paragraph.add_run(header_text)
                header_run.font.color.rgb = RGBColor(128, 128, 128)
                header_run.font.size = Pt(11)

            doc.add_paragraph()

            # Divide o texto em linhas e adiciona ao documento
            linhas_recibo = texto_formatado.split('\n')
            for linha in linhas_recibo:
                if linha.strip():
                    p = doc.add_paragraph()
                    if "RECIBO Nº" in linha and "VALOR" in linha:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        partes = linha.split("VALOR:")
                        run = p.add_run(partes[0].strip())
                        run.bold = True
                        p.add_run('\t')
                        tab_stop = p.paragraph_format.tab_stops.add_tab_stop(
                            Inches(6),
                            WD_ALIGN_PARAGRAPH.RIGHT
                        )
                        valor_texto = f"VALOR:{partes[1].strip()}"
                        run = p.add_run(valor_texto)
                        run.bold = True
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.add_run(linha.strip())

            doc.add_paragraph()

            # Data em português usando a data selecionada
            mes_pt = traduzir_mes(data_selecionada.strftime('%B'))
            data_formatada_completa = f"Bauru, {data_selecionada.day} de {mes_pt} de {data_selecionada.year}"
            
            data_paragraph = doc.add_paragraph()
            data_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            data_paragraph.add_run(data_formatada_completa)

            doc.add_paragraph()
            
            # Assinatura
            assinatura_paragraph = doc.add_paragraph()
            assinatura_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            assinatura_paragraph.add_run("_" * 50 + "\n")
            assinatura_paragraph.add_run(cliente.razao_social.upper() + "\n")
            assinatura_paragraph.add_run(cliente.cpf_cnpj)

            # Salvar documento
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Salvar no banco com a data selecionada
            recibo = ReciboGerado(
                numero_recibo=numero_recibo,
                modelo_id=int(modelo_id),
                cliente_nome=cliente_nome,
                valor=valor_float,
                data_geracao=data_selecionada,  # Salva a data selecionada
                documento_blob=doc_buffer.getvalue()
            )
            db.session.add(recibo)
            db.session.commit()

            documentos_gerados.append((cliente_nome, doc_buffer.getvalue()))
            preview_content.append({
                'id': recibo.id,
                'nome': cliente_nome,
                'conteudo': linhas_recibo
            })

        return jsonify({
            'preview': preview_content,
            'status': 'success'
        })

    except Exception as e:
        print(f"Erro detalhado: {str(e)}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500



@app.route('/download_recibos', methods=['GET', 'POST'])
def download_recibos():
    try:
        if request.method == 'POST':
            # Se for POST, pega os IDs dos recibos enviados
            dados = request.json
            recibos_ids = dados.get('recibos', [])
            recibos = ReciboGerado.query.filter(ReciboGerado.id.in_(recibos_ids)).all()
        else:
            # Se for GET, usa a variável global (compatibilidade com versão anterior)
            global documentos_gerados
            if not documentos_gerados:
                return jsonify({'error': 'Nenhum documento disponível para download'}), 404
            recibos = None
        
        # Cria um buffer em memória para o ZIP
        zip_buffer = io.BytesIO()
        
        # Cria o arquivo ZIP com modo de escrita binária e compressão
        with zipfile.ZipFile(zip_buffer, 'w', compression=zipfile.ZIP_DEFLATED) as zip_file:
            if recibos:
                # Caso esteja usando recibos do banco de dados (POST)
                for recibo in recibos:
                    doc_buffer = io.BytesIO(recibo.documento_blob)
                    doc_buffer.seek(0)
                    
                    nome_arquivo = "".join(c for c in recibo.cliente_nome if c.isalnum() or c in (' ', '-', '_'))
                    zip_file.writestr(f"recibo_{recibo.numero_recibo}_{nome_arquivo}.docx", doc_buffer.getvalue())
                    doc_buffer.close()
            else:
                # Caso esteja usando a variável global (GET)
                for nome, doc_content in documentos_gerados:
                    doc_buffer = io.BytesIO(doc_content)
                    doc_buffer.seek(0)
                    
                    nome_arquivo = "".join(c for c in nome if c.isalnum() or c in (' ', '-', '_'))
                    zip_file.writestr(f"recibo_{nome_arquivo}.docx", doc_buffer.getvalue())
                    doc_buffer.close()
        
        # Prepara o buffer para leitura
        zip_buffer.seek(0)
        zip_size = zip_buffer.getbuffer().nbytes
        
        print(f"ZIP gerado com sucesso. Tamanho: {zip_size} bytes")
        
        # Envia o arquivo
        response = send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name='recibos.zip'
        )
        
        # Headers específicos
        response.headers["Content-Length"] = zip_size
        response.headers["Content-Type"] = "application/zip"
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        
        return response

    except Exception as e:
        print(f"Erro no download: {str(e)}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500



@app.route('/historico_recibos')
def historico_recibos():
    recibos = ReciboGerado.query.order_by(ReciboGerado.data_geracao.desc()).all()
    return render_template('historico.html', recibos=recibos)

@app.route('/consulta_recibos')
def consulta_recibos():
    recibos = ReciboGerado.query.order_by(ReciboGerado.data_geracao.desc()).all()
    return render_template('consulta_recibos.html', recibos=recibos)

@app.route('/download_recibo/<int:recibo_id>')
def download_recibo(recibo_id):
    recibo = ReciboGerado.query.get_or_404(recibo_id)
    
    return send_file(
        io.BytesIO(recibo.documento_blob),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'recibo_{recibo.numero_recibo}.docx'
    )

@app.route('/visualizar_recibo/<int:recibo_id>')
def visualizar_recibo(recibo_id):
    recibo = ReciboGerado.query.get_or_404(recibo_id)
    data_formatada = recibo.data_geracao.strftime('%d/%m/%Y')
    mes_pt = traduzir_mes(recibo.data_geracao.strftime('%B'))
    data_completa = f"Bauru, {recibo.data_geracao.day} de {mes_pt} de {recibo.data_geracao.year}"
    return render_template('visualizar_recibo.html', recibo=recibo, data_formatada=data_completa)



@app.route('/modelos', methods=['GET'])
def listar_modelos():
    modelos = ModeloRecibo.query.all()
    return jsonify([{
        'id': m.id,
        'nome': m.nome,
        'conteudo': m.conteudo,
        'header_text': m.header_text,
        'logo_path': m.logo_path.replace('\\', '/') if m.logo_path else None
    } for m in modelos])

@app.route('/salvar_modelo', methods=['POST'])
def salvar_modelo():
    try:
        dados = request.json
        modelo_id = dados.get('modelo_id')
        nome = dados.get('nome')
        conteudo = dados.get('conteudo')
        header_text = dados.get('header_text')
        logo_path = dados.get('logo_path')
        
        modelo = ModeloRecibo.query.get(modelo_id)
        if not modelo:
            modelo = ModeloRecibo(id=modelo_id)
            db.session.add(modelo)
            
        modelo.nome = nome
        modelo.conteudo = conteudo
        modelo.header_text = header_text
        if logo_path:
            modelo.logo_path = logo_path
            
        db.session.commit()
        
        return jsonify({
            'status': 'sucesso',
            'modelo': {
                'id': modelo.id,
                'nome': modelo.nome,
                'conteudo': modelo.conteudo,
                'header_text': modelo.header_text,
                'logo_path': modelo.logo_path
            }
        })
        
    except Exception as e:
        return jsonify({'erro': str(e)}), 500


@app.route('/modelos/<int:modelo_id>', methods=['PUT'])
def atualizar_modelo(modelo_id):
    dados = request.json
    modelo = ModeloRecibo.query.get_or_404(modelo_id)
    modelo.nome = dados['nome']
    modelo.conteudo = dados['conteudo']
    db.session.commit()
    return jsonify({'message': 'Modelo atualizado com sucesso'})

@app.route('/add_cliente', methods=['POST'])
def add_cliente():
    try:
        data = request.json
        documento = data['cpf_cnpj']
        documento_limpo = ''.join(filter(str.isdigit, documento))
        
        # Determina e valida o tipo de documento
        if len(documento_limpo) <= 11:
            if not validar_cpf(documento_limpo):
                return jsonify({'error': 'CPF inválido'}), 400
            tipo = 'pessoa'
        else:
            if not validar_cnpj(documento_limpo):
                return jsonify({'error': 'CNPJ inválido'}), 400
            tipo = 'empresa'
            
        # Verifica se já existe cliente com este documento
        cliente_existente = Cliente.query.filter_by(cpf_cnpj=documento).first()
        if cliente_existente:
            return jsonify({'error': 'CPF/CNPJ já cadastrado'}), 400
            
        novo_cliente = Cliente(
            razao_social=data['razao_social'],
            cpf_cnpj=documento,
            tipo=tipo
        )
        
        db.session.add(novo_cliente)
        db.session.commit()
        
        return jsonify({'message': 'Cliente cadastrado com sucesso'}), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def redimensionar_logo(arquivo, tamanho_max=(300, 100)):
    img = Image.open(arquivo)
    img.thumbnail(tamanho_max)
    return img


@app.route('/upload_logo', methods=['POST'])
def upload_logo():
    try:
        file = request.files['logo']
        modelo_id = request.form.get('modelo_id')
        
        # Criar pasta logos se não existir
        logos_dir = 'static/images/logos'
        os.makedirs(logos_dir, exist_ok=True)
        
        # Gerar nome único para o arquivo
        filename = f'logo_modelo_{modelo_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.png'
        filepath = f'/{logos_dir}/{filename}'  # Adiciona a barra no início
        
        # Remove a barra inicial para salvar o arquivo
        file.save(filepath[1:])  # Remove a primeira barra para salvar
        
        # Atualizar modelo com o caminho começando com /
        modelo = ModeloRecibo.query.get(modelo_id)
        if modelo:
            modelo.logo_path = filepath  # Salva com a barra no início
            db.session.commit()
        
        return jsonify({
            'success': True,
            'logo_path': filepath  # Retorna com a barra no início
        })
        
    except Exception as e:
        print(f"Erro no upload: {str(e)}")
        return jsonify({'error': str(e)}), 500
    
@app.route('/consulta_clientes')
def consulta_clientes():
    clientes = Cliente.query.order_by(Cliente.razao_social).all()
    return render_template('consulta_clientes.html', clientes=clientes)

@app.route('/delete_cliente/<int:cliente_id>', methods=['DELETE'])
def delete_cliente(cliente_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        db.session.delete(cliente)
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)})
    
@app.route('/atualizar_recibo', methods=['POST'])
def atualizar_recibo():
    print("Iniciando atualização do recibo")
    try:
        dados = request.json
        recibo_id = dados.get('recibo_id')
        conteudo_novo = dados.get('conteudo')

        print(f"Recibo ID: {recibo_id}")
        print(f"Novo conteúdo recebido: {conteudo_novo}")

        # Busca o recibo e o modelo associado
        recibo = ReciboGerado.query.get_or_404(recibo_id)
        modelo = ModeloRecibo.query.get(recibo.modelo_id)  # Obtém o modelo associado ao recibo
        
        print(f"Recibo encontrado: {recibo.numero_recibo}")
        print(f"Modelo associado: {modelo.id if modelo else 'Nenhum'}")

        # Criar um novo documento Word        
        doc = Document()
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Cabeçalho com logo e informações
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(1.2)
        header_table.columns[1].width = Inches(5.8)

        # Logo
        logo_cell = header_table.cell(0, 0)
        logo_paragraph = logo_cell.paragraphs[0]
        logo_run = logo_paragraph.add_run()

        if modelo and modelo.logo_path:
            logo_path = modelo.logo_path.lstrip('/')  # Remove a barra inicial se existir
            if os.path.exists(logo_path):
                try:
                    logo_run.add_picture(logo_path, width=Inches(1.2))
                except Exception as e:
                    print(f"Erro ao adicionar logo personalizada: {str(e)}")
                    logo_run.add_picture('static/images/logo.png', width=Inches(1.2))
            else:
                logo_run.add_picture('static/images/logo.png', width=Inches(1.2))
        else:
            logo_run.add_picture('static/images/logo.png', width=Inches(1.2))

        # Texto do cabeçalho
        info_cell = header_table.cell(0, 1)
        info_paragraph = info_cell.paragraphs[0]
        header_text = modelo.header_text if (modelo and modelo.header_text) else "BEIJO E MATOS CONSTRUÇÕES E ENGENHARIA LTDA\nJoaquim da Silva Martha, 12-53 - Sala 3 - Altos da Cidade - Bauru/SP\nguilhermebeijo@bencato.com.br - CNPJ: 26.149.105/0001-09 - www.bencato.com.br"
        info_run = info_paragraph.add_run(header_text)
        info_run.font.color.rgb = RGBColor(128, 128, 128)
        info_run.font.size = Pt(11)

        doc.add_paragraph()  # Espaço após cabeçalho

        # Adiciona conteúdo atualizado com formatação específica para a linha do recibo
        for linha in conteudo_novo:
            if linha.strip():
                p = doc.add_paragraph()               
                if "RECIBO Nº" in linha and "VALOR" in linha:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alinha o parágrafo à esquerda
                    
                    # Calcula o espaço necessário para alinhar à direita
                    espaco_total = section.page_width - section.left_margin - section.right_margin
                    
                    # Divide a linha em duas partes: número do recibo e valor
                    partes = linha.split("VALOR:")
                    
                    # Adiciona a primeira parte (RECIBO Nº) em negrito
                    run = p.add_run(partes[0].strip())
                    run.bold = True
                    
                    # Adiciona tabulação para alinhar à direita
                    p.add_run('\t')  # Adiciona uma tabulação
                    
                    # Configura a tabulação para alinhar à direita
                    tab_stop = p.paragraph_format.tab_stops.add_tab_stop(
                        Inches(6),  # Posição da tabulação (ajuste conforme necessário)
                        WD_ALIGN_PARAGRAPH.RIGHT
                    )
                    
                    # Adiciona "VALOR:" e o valor alinhado à direita
                    valor_texto = f"VALOR:{partes[1].strip()}"
                    run = p.add_run(valor_texto)
                    run.bold = True
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.add_run(linha.strip())

        doc.add_paragraph()  # Espaço antes da data

        # Dentro da função atualizar_recibo
        data_recebida = dados.get('data')
        if data_recebida:
            data_selecionada = datetime.strptime(data_recebida, '%Y-%m-%d')
            mes_pt = traduzir_mes(data_selecionada.strftime('%B'))
            data_formatada = f"Bauru, {data_selecionada.day} de {mes_pt} de {data_selecionada.year}"
        else:
            # Usa a data do recibo existente
            data_selecionada = recibo.data_geracao
            mes_pt = traduzir_mes(data_selecionada.strftime('%B'))
            data_formatada = f"Bauru, {data_selecionada.day} de {mes_pt} de {data_selecionada.year}"

        # Na parte de geração do documento
        data_paragraph = doc.add_paragraph()
        data_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        data_paragraph.add_run(data_formatada)  # Usar a data formatada que definimos


        doc.add_paragraph()  # Espaço antes da assinatura

        # Busca cliente para informações da assinatura
        cliente = Cliente.query.filter_by(razao_social=recibo.cliente_nome).first()

        # Adiciona assinatura
        assinatura_paragraph = doc.add_paragraph()
        assinatura_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        assinatura_paragraph.add_run("_" * 50 + "\n")
        assinatura_paragraph.add_run(cliente.razao_social.upper() + "\n")
        assinatura_paragraph.add_run(cliente.cpf_cnpj)

        # Salva o documento
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        # Atualiza no banco de dados
        recibo.documento_blob = doc_buffer.getvalue()
        db.session.commit()

        print("Recibo atualizado com sucesso no banco de dados")

        return jsonify({'status': 'sucesso', 'mensagem': 'Recibo atualizado com sucesso'})

    except Exception as e:
        print(f"Erro ao atualizar recibo: {str(e)}")
        import traceback
        print(traceback.format_exc())
        return jsonify({'status': 'erro', 'mensagem': str(e)}), 500

@app.route('/debug_recibo/<int:recibo_id>')
def debug_recibo(recibo_id):
    try:
        print(f"Buscando recibo ID: {recibo_id}")
        recibo = ReciboGerado.query.get(recibo_id)
        
        if not recibo:
            print(f"Recibo {recibo_id} não encontrado")
            return jsonify({'erro': f'Recibo {recibo_id} não encontrado'}), 404
            
        print(f"Recibo encontrado: {recibo.numero_recibo}")
        
        # Lê o documento Word
        doc = Document(io.BytesIO(recibo.documento_blob))
        
        # Extrai conteúdo
        conteudo = []
        for paragrafo in doc.paragraphs:
            texto = paragrafo.text.strip()
            if texto:
                conteudo.append(texto)
                print(f"Parágrafo encontrado: {texto}")
        
        dados = {
            'id': recibo.id,
            'numero_recibo': recibo.numero_recibo,
            'cliente_nome': recibo.cliente_nome,
            'valor': str(recibo.valor),
            'data_geracao': recibo.data_geracao.strftime('%d/%m/%Y %H:%M'),
            'conteudo': conteudo
        }
        
        print("Dados do recibo:", json.dumps(dados, indent=2))
        return jsonify(dados)
        
    except Exception as e:
        print(f"Erro ao debugar recibo: {str(e)}")
        return jsonify({'erro': str(e)}), 500

@app.route('/debug_modelos')
def debug_modelos():
    modelos = ModeloRecibo.query.all()
    return jsonify([{
        'id': m.id,
        'nome': m.nome,
        'conteudo': m.conteudo
    } for m in modelos])

@app.route('/debug_modelo/<int:modelo_id>')
def debug_modelo(modelo_id):
    modelo = ModeloRecibo.query.get(modelo_id)
    if modelo:
        return jsonify({
            'id': modelo.id,
            'nome': modelo.nome,
            'logo_path': modelo.logo_path,
            'header_text': modelo.header_text
        })
    return jsonify({'error': 'Modelo não encontrado'}), 404

@app.route('/reset_database', methods=['POST'])
def reset_database():
    try:
        with app.app_context():
            # Deletar todos os recibos
            ReciboGerado.query.delete()
            
            # Resetar o contador
            seq = ReceiptSequence.query.first()
            if seq:
                seq.last_number = 0
            else:
                seq = ReceiptSequence(last_number=0)
                db.session.add(seq)
                
            db.session.commit()
        return jsonify({'message': 'Database resetada com sucesso'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def init_db():    
    with app.app_context():        # Criar tabelas se não existirem
        db.create_all()        
        # Criar modelos padrão se não existirem
        if not ModeloRecibo.query.first():
            modelos_padrao = [
                {
                    'nome': 'Modelo Emitente',
                    'conteudo': 'RECIBO Nº {numero_recibo}...'
                },
                {
                    'nome': 'Modelo Destinatário',
                    'conteudo': 'RECIBO Nº {numero_recibo}...'
                },
                {
                    'nome': 'Modelo Personalizado',
                    'conteudo': 'RECIBO Nº {numero_recibo}...'
                }
            ]
            for modelo in modelos_padrao:
                db.session.add(ModeloRecibo(**modelo))
            db.session.commit()

if __name__ == '__main__':
    with app.app_context():
        init_db()
    app.run(debug=True)
